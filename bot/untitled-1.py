# -*- coding: utf-8 -*-
import telebot
from random import *
from telebot import types
import docx

with open("facts.txt", "r", encoding="utf8") as f:
    lines = f.readlines()
with open("links.txt", "r", encoding="utf8") as f1:
    lines1 = f1.readlines()

photo1 = open("child.jpg", "rb")
photo3 = open("war.jpg", "rb")
photo4 = open("death.jpg", "rb")
audio = open('audio.m4a', 'rb')
doc = docx.Document('child.docx')
doc1 = docx.Document('war.docx')
doc2 = docx.Document('death.docx')
lines2 = [p.text for p in doc.paragraphs]
lines3 = [p.text for p in doc1.paragraphs]
lines4 = [p.text for p in doc2.paragraphs]
bot = telebot.TeleBot('6228487978:AAElkCgqSAz5HFuWoFin9jDitKQSdssT_9g')


@bot.message_handler(commands=["start"])
def start(message):
    markup = types.ReplyKeyboardMarkup(resize_keyboard=True)
    btn1 = types.KeyboardButton("Информация о боте")
    btn2 = types.KeyboardButton("Узнать случайный факт о Ватутине")
    btn3 = types.KeyboardButton("Биография Ватутина по периодам жизни")
    btn4 = types.KeyboardButton("Полезные ссылки на статьи и фильмы о Ватутине")
    btn5 = types.KeyboardButton("Аудиоэкскурсия по музею Ватутина")
    markup.add(btn1, btn2, btn3, btn4, btn5)
    bot.send_message(
        message.chat.id,
        text="Привет, {0.first_name}! Я - бот, который поможет тебе больше узнать о биографии великого полководца Николая Фёдоровича Ватутина. В боте можно ориентироваться с помощью кнопок.".format(
            message.from_user
        ),
        reply_markup=markup,
    )


@bot.message_handler(content_types=["text"])
def func(message):
    if message.text == "Информация о боте":
        bot.send_message(
            message.chat.id,
            text="Бот создан для облегчения патриотического воспитания обучающихся с помощью ИКТ.\nАвтор бота - Козлов Дмитрий Иванович, Белгород, 2023",
        )
    elif message.text == "Узнать случайный факт о Ватутине":
        markup = types.ReplyKeyboardMarkup(resize_keyboard=True)
        back = types.KeyboardButton("Вернуться в главное меню")
        markup.add(back)
        bot.send_message(
            message.chat.id, text=lines[randint(0, len(lines) - 1)], reply_markup=markup
        )
    elif message.text == "Аудиоэкскурсия по музею Ватутина":
        markup = types.ReplyKeyboardMarkup(resize_keyboard=True)
        back = types.KeyboardButton("Вернуться в главное меню")
        markup.add(back)
        bot.send_audio(message.chat.id, audio=audio, caption='Аудиоэкскурсия по дому-музею Ватутина', reply_markup=markup)
    elif message.text == "Биография Ватутина по периодам жизни":
        markup = types.ReplyKeyboardMarkup(resize_keyboard=True)
        btn1 = types.KeyboardButton("Детство и юность")
        btn2 = types.KeyboardButton("Военная карьера")
        btn3 = types.KeyboardButton("Гибель полководца")
        back = types.KeyboardButton("Вернуться в главное меню")
        markup.add(btn1, btn2, btn3, back)
        bot.send_message(
            message.chat.id, text="Выберите нужный период:", reply_markup=markup
        )
    elif message.text == "Детство и юность":
        bot.send_photo(message.chat.id, photo=photo1)
        text = ''.join([line.strip() for line in lines2])
        while text:
            bot.send_message(message.chat.id, text[:4090])
            text = text[4090:]
    elif message.text == "Военная карьера":
        bot.send_photo(message.chat.id, photo=photo3)
        text = ''.join([line.strip() for line in lines3])
        while text:
            bot.send_message(message.chat.id, text[:4090])
            text = text[4090:]
    elif message.text == "Гибель полководца":
        bot.send_photo(message.chat.id, photo=photo4)
        text = ''.join([line.strip() for line in lines4])
        while text:
            bot.send_message(message.chat.id, text[:4090])
            text = text[4090:]
    elif message.text == "Вернуться в главное меню":
        markup = types.ReplyKeyboardMarkup(resize_keyboard=True)
        btn1 = types.KeyboardButton("Информация о боте")
        btn2 = types.KeyboardButton("Узнать случайный факт о Ватутине")
        btn3 = types.KeyboardButton("Биография Ватутина по периодам жизни")
        btn4 = types.KeyboardButton("Полезные ссылки на статьи и фильмы о Ватутине")
        btn5 = types.KeyboardButton("Аудиоэкскурсия по музею Ватутина")
        markup.add(btn1, btn2, btn3, btn4, btn5)
        bot.send_message(
            message.chat.id, text="Вы вернулись в главное меню", reply_markup=markup
        )
    elif message.text == "Полезные ссылки на статьи и фильмы о Ватутине":
        markup = types.ReplyKeyboardMarkup(resize_keyboard=True)
        back = types.KeyboardButton("Вернуться в главное меню")
        markup.add(back)
        bot.send_message(message.chat.id, text="\n".join(lines1), reply_markup=markup)

    elif message.text == "Вернуться в главное меню":
        markup = types.ReplyKeyboardMarkup(resize_keyboard=True)
        btn1 = types.KeyboardButton("Информация о боте")
        btn2 = types.KeyboardButton("Узнать случайный факт о Ватутине")
        btn3 = types.KeyboardButton("Биография Ватутина по периодам жизни")
        btn4 = types.KeyboardButton("Полезные ссылки на статьи и фильмы о Ватутине")
        btn5 = types.KeyboardButton("Аудиоэкскурсия по музею Ватутина")
        markup.add(btn1, btn2, btn3, btn4, btn5)
        bot.send_message(
            message.chat.id, text="Вы вернулись в главное меню", reply_markup=markup
        )
    elif message.text != 'Биография Ватутина по периодам жизни':
        bot.send_message(
            message.chat.id,
            text="Вы ввели что-то неправильное! Пожалуйста, воспользуйтесь кнопками!",
        )
    else:
        bot.send_message(
            message.chat.id,
            text="Вы ввели что-то неправильное! Пожалуйста, воспользуйтесь кнопками!",
        )


bot.polling(none_stop=True)
